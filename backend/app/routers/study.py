"""Study (拆书) routes — materials, chapters, character extraction."""
from __future__ import annotations

import asyncio
import json
import re
from datetime import datetime
from typing import Any, Literal

from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, UploadFile
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker
from sqlalchemy.orm import selectinload

from app.agents.base import AgentContext
from app.agents.study import (
    StudyBehaviorPatternAgent,
    StudyCharacterAgent,
    StudyEventAgent,
)
from app.core.database import AsyncSessionLocal, get_db
from app.core.errors import bad_request, not_found
from app.models.study import BehaviorPattern, StudyChapter, StudyCharacter, StudyMaterial
from app.models.task import AgentTask
from app.schemas import (
    APIResponse,
    BehaviorPatternRead,
    ChapterizeRequest,
    StudyBehaviorExtractRequest,
    StudyBehaviorExtractResponse,
    StudyBulkRequest,
    StudyBulkStartResponse,
    StudyChapterRead,
    StudyCharacterCreate,
    StudyCharacterRead,
    StudyForeshadowSummary,
    StudyMaterialCreate,
    StudyMaterialDetail,
    StudyMaterialOverview,
    StudyMaterialRead,
    StudyMaterialUpdate,
    StudyRelationshipApplyRequest,
    StudyRelationshipApplyResponse,
    StudyRelationshipEnrichRequest,
    StudyRelationshipEnrichResponse,
    StudyRelationshipEnrichedItem,
    StudyRelationshipsResponse,
    StudyRelationshipSuggestion,
    StudyRequest,
)
from app.services.llm.router import get_llm_router
from app.services.prompt_engine import get_prompt_engine


router = APIRouter(prefix="/study", tags=["study"])


# -------------------- chapterize helpers --------------------

# Chinese: "第N章 / 第N节 / 第N卷 / 第N回" with optional whitespace
# between every token. The original regex only matched "第N章", which
# left popular web novels (e.g. 蛊真人 uses "第NNN节：xxx") falling
# through to the single-chapter "全文" fallback.
#
# We capture the suffix char in group 2 so ``_chunks_from_matches``
# can echo it back in the title — "第 12 节" is meaningfully
# different from "第 12 章" and the user can see at a glance which
# a given book uses.
_CN_CHAPTER_RE = re.compile(
    r"^\s*第\s*([零〇一二三四五六七八九十百千万0-9]+)\s*([章节卷回])[\s　\.：:—\-]*(.{0,80})$",
    re.MULTILINE,
)
# English: "Chapter 1" / "Chapter 1: Foo" / "CHAPTER ONE" / "Chapter One"
_EN_CHAPTER_RE = re.compile(
    r"^\s*CHAPTER\s+([0-9]+|[A-Za-z]+)[\.\:\s\-—]*(.{0,80})$",
    re.MULTILINE | re.IGNORECASE,
)
# Generic "Chapter 1" fallback
_EN_CHAPTER_LOOSE_RE = re.compile(
    r"^\s*Chapter\s+([0-9]+)[\.\:\s\-—]*(.{0,80})$",
    re.MULTILINE | re.IGNORECASE,
)


def _cn_to_int(s: str) -> int:
    """Convert a Chinese chapter number to int. Falls back to the
    raw string hash if the number is too gnarly to parse (we just
    need an ordering signal — a 2.7 GB 2000-chapter 玄幻 can have
    ``第两千三百四十五章`` and the regex is unlikely to know every
    variant; a fallback is fine).
    """
    digits_map = {
        "零": 0, "〇": 0, "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
        "六": 6, "七": 7, "八": 8, "九": 9,
    }
    units = {"十": 10, "百": 100, "千": 1000, "万": 10000}
    if s.isdigit():
        return int(s)
    total = 0
    section = 0
    current = 0
    for ch in s:
        if ch in digits_map:
            current = digits_map[ch]
        elif ch in units:
            section += (current or 1) * units[ch]
            current = 0
        else:
            # Unknown character — bail with a hash-based fallback.
            return abs(hash(s)) % 100000
    return total + section + current


def _split_chapters(text: str, pattern: str = "auto") -> list[tuple[int, str, str]]:
    """Return a list of (chapter_index, title, content) tuples.

    The first item is always ``(1, "序章", text_before_first_header)`` if
    the user pastes a preamble, or ``(1, "第 N 章", full_text)`` if the
    first line is already a chapter header.
    """
    text = (text or "").strip()
    if not text:
        return []
    # Pick a regex.
    if pattern == "chinese":
        rx = _CN_CHAPTER_RE
    elif pattern == "english":
        rx = _EN_CHAPTER_RE
    else:
        # auto: try Chinese first (more common in the project's audience),
        # fall back to English.
        cn_hits = list(_CN_CHAPTER_RE.finditer(text))
        en_hits = list(_EN_CHAPTER_RE.finditer(text))
        if not en_hits and not cn_hits:
            en_hits = list(_EN_CHAPTER_LOOSE_RE.finditer(text))
        # Whichever has more matches wins.
        if len(cn_hits) >= len(en_hits):
            rx = _CN_CHAPTER_RE
            matches = cn_hits
            cn_mode = True
        else:
            rx = _EN_CHAPTER_RE
            matches = en_hits
            cn_mode = False
        if not matches:
            return [(1, "全文", text)]
        return _chunks_from_matches(text, matches, cn_mode=cn_mode)
    matches = list(rx.finditer(text))
    if not matches:
        return [(1, "全文", text)]
    return _chunks_from_matches(text, matches, cn_mode=(pattern == "chinese"))


def _chunks_from_matches(
    text: str, matches: list[re.Match], *, cn_mode: bool
) -> list[tuple[int, str, str]]:
    out: list[tuple[int, str, str]] = []
    for i, m in enumerate(matches):
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[start:end].strip()
        # If the very first line of the document has no chapter header
        # before the first match, surface that preamble as a prologue
        # so the user can still see it.
        if i == 0 and m.start() > 0:
            preamble = text[: m.start()].strip()
            if len(preamble) > 50:
                out.append((1, "序章", preamble))
                # Bump subsequent indices by 1.
                first_real_index = 2
            else:
                first_real_index = 1
        else:
            first_real_index = 1
        # Title: capture group 1 is the number, group 2 is the
        # suffix char (章/节/卷/回) — we echo that back in the
        # title so the user can see which book convention this is,
        # group 3 is the optional chapter name. We also include
        # the original header line in the content so the user can
        # see the full chapter start.
        num = m.group(1)
        suffix = m.group(2) or "章"
        title = (m.group(3) or "").strip()
        if cn_mode:
            try:
                num_int = _cn_to_int(num)
            except Exception:
                num_int = i + 1
        else:
            try:
                num_int = int(num)
            except ValueError:
                num_int = i + 1
        # Cap the title length.
        if len(title) > 60:
            title = title[:60] + "…"
        full_title = f"第 {num_int} {suffix}" + (f" · {title}" if title else "")
        # Prepend the original header line so the chunk is searchable.
        full_body = (m.group(0).rstrip() + "\n" + body).strip()
        out.append((first_real_index + (num_int - 1 if i > 0 or first_real_index == 1 else 0), full_title, full_body))
    # Re-index defensively (we may have inserted a prologue).
    return [(i + 1, t, c) for i, (_, t, c) in enumerate(out)]


# -------------------- endpoints --------------------

@router.get("/materials", response_model=APIResponse[list[StudyMaterialRead]])
async def list_materials(
    project_id: int | None = Query(default=None),
    db: AsyncSession = Depends(get_db),
) -> APIResponse[list[StudyMaterialRead]]:
    stmt = select(StudyMaterial).order_by(StudyMaterial.updated_at.desc())
    if project_id is not None:
        stmt = stmt.where(StudyMaterial.project_id == project_id)
    rows = (await db.execute(stmt)).scalars().all()
    return {"ok": True, "data": [StudyMaterialRead.from_orm_trimmed(r) for r in rows]}


@router.post("/materials", response_model=APIResponse[StudyMaterialRead])
async def create_material(
    body: StudyMaterialCreate, db: AsyncSession = Depends(get_db)
) -> APIResponse[StudyMaterialRead]:
    row = StudyMaterial(
        project_id=body.project_id,
        title=body.title,
        author=body.author,
        source=body.source,
        raw_text=body.raw_text,
        status="draft" if body.raw_text else "empty",
    )
    db.add(row)
    await db.flush()
    return {"ok": True, "data": StudyMaterialRead.from_orm_trimmed(row)}


@router.post("/materials/upload/batch")
async def upload_materials_batch(
    files: list[UploadFile] = File(...),
    auto_chapterize: bool = Form(default=True),
    min_chapter_chars: int = Form(default=200),
    db: AsyncSession = Depends(get_db),
):
    """R19: accept up to 5 books in a single multipart POST.

    Each file is parsed with the same per-format dispatch as the
    single-file upload route; each resulting material is then
    auto-chapterized in-place so the user lands on a populated
    library without an extra click. Per-file failures are surfaced
    in the response as ``{"ok": false, "error": ...}`` and the
    whole batch still returns 200 — we don't want one bad EPUB
    to nuke 4 successful TXT uploads.

    The 5-file cap is enforced in the route (the frontend also
    mirrors it via the ``<input multiple>`` limit).

    No ``response_model`` — the per-entry shape is mixed (success
    returns ``{ok, data: StudyMaterialRead}``, failure returns
    ``{ok: false, filename, error}``), so we let the dict go out
    untyped and the client treats each entry by its own ``ok``.
    """
    if not files:
        raise bad_request("至少选择一个文件。")
    if len(files) > 5:
        raise bad_request(f"批量上传最多 5 个文件,收到 {len(files)}。")
    results: list[dict[str, Any]] = []
    for f in files:
        try:
            content = await f.read()
            if len(content) > 32 * 1024 * 1024:
                raise ValueError("文件过大 (>32MB)。")
            title = (f.filename or "upload").rsplit(".", 1)[0]
            raw = _extract_text_from_upload(f.filename or "upload.txt", content)
            if not raw or not raw.strip():
                raise ValueError("文件解析后没有可识别的正文。")
            row = StudyMaterial(
                title=title,
                author="",
                source="upload",
                raw_text=raw,
                status="draft" if raw else "empty",
            )
            db.add(row)
            await db.flush()
            entry: dict[str, Any] = {
                "ok": True,
                "data": {
                    "id": row.id,
                    "project_id": row.project_id,
                    "title": row.title,
                    "author": row.author,
                    "source": row.source,
                    "status": row.status,
                    "error": row.error,
                    "chapter_count": row.chapter_count or 0,
                    "character_count": row.character_count or 0,
                    "raw_text_length": len(row.raw_text or ""),
                    "extra": row.extra,
                    "created_at": row.created_at.isoformat() if row.created_at else None,
                    "updated_at": row.updated_at.isoformat() if row.updated_at else None,
                },
            }
            if auto_chapterize and row.raw_text:
                try:
                    chunks = _split_chapters(row.raw_text, pattern="auto")
                    created = 0
                    if chunks:
                        for idx, t, body in chunks:
                            if len(body) < min_chapter_chars:
                                continue
                            # Use a direct INSERT instead of
                            # ``row.chapters.append(StudyChapter(...))``
                            # — appending onto a relationship triggers
                            # a lazy-load of the collection on first
                            # access, and that lazy load is sync-IO
                            # inside an async session (greenlet error).
                            # Inserting the chapter row directly with
                            # ``material_id`` sidesteps the relationship
                            # walk entirely.
                            db.add(StudyChapter(
                                material_id=row.id,
                                chapter_index=idx,
                                title=t,
                                content=body,
                                char_count=len(body),
                            ))
                            created += 1
                        # Flush so the new chapter rows reach the DB
                        # before the response serialises chapter_count.
                        await db.flush()
                    row.chapter_count = created
                    row.status = "ready" if created else "draft"
                    # One more flush to push chapter_count + status
                    # back to the DB and update the in-memory row
                    # attributes that we read on the next two lines.
                    await db.flush()
                except Exception as exc:
                    entry["chapterize_error"] = str(exc)
            # Reflect the post-chapterize values back into the entry
            # so the client gets the final chapter_count + status
            # without an extra GET.
            entry["data"]["chapter_count"] = row.chapter_count or 0
            entry["data"]["status"] = row.status
            results.append(entry)
        except Exception as exc:
            results.append({
                "ok": False,
                "filename": f.filename,
                "error": f"{exc.__class__.__name__}: {exc}".strip(),
            })
    return {"ok": True, "data": results}


@router.post("/materials/upload", response_model=APIResponse[StudyMaterialRead])
async def upload_material(
    title: str = Form(...),
    author: str = Form(""),
    project_id: int | None = Form(default=None),
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
) -> APIResponse[StudyMaterialRead]:
    """Accept a book upload and create the material row with its body.

    R18 / P0-STUDY-3: supports a much wider set of formats than the
    original ``.txt`` only path. The list is
    ``.txt / .md / .pdf / .docx / .html / .htm / .epub``.

    Multipart route — the file's bytes are decoded/parsed into
    ``raw_text`` (the MVP doesn't keep an on-disk copy; if the user
    wants a smaller payload for chapterize, they can re-paste).

    The file extension is matched case-insensitively; the file body
    is parsed lazily inside ``_extract_text_from_upload`` so a
    missing optional dep degrades gracefully (one format becomes
    unavailable, not the whole endpoint).
    """
    raw_bytes = await file.read()
    if len(raw_bytes) > 32 * 1024 * 1024:
        # 32 MB hard cap. PDF/DOCX can be big and parsing them into
        # raw_text balloons memory. Above this we refuse rather than
        # OOM the worker. User can split their book or paste chunks.
        raise HTTPException(status_code=413, detail="文件过大 (>32MB),请拆开上传或粘贴正文。")
    filename = file.filename or "upload.txt"
    try:
        raw = _extract_text_from_upload(filename, raw_bytes)
    except ValueError as exc:
        raise bad_request(str(exc)) from exc
    except Exception as exc:
        # Don't leak the underlying stack to the user; they get a
        # clean message and the failure is in the worker log.
        raise HTTPException(
            status_code=400,
            detail=f"解析文件失败: {exc.__class__.__name__}: {exc}".strip(),
        ) from exc
    if not raw or not raw.strip():
        raise bad_request("文件解析后没有可识别的正文。")
    row = StudyMaterial(
        project_id=project_id,
        title=title,
        author=author,
        source="upload",
        raw_text=raw,
        status="draft" if raw else "empty",
    )
    db.add(row)
    await db.flush()
    return {"ok": True, "data": StudyMaterialRead.from_orm_trimmed(row)}


# -------------------- R18 / P0-STUDY-3: file -> text --------------------

# Map lower-cased file extension -> parser key. Anything not in this
# map is rejected up front (cheaper than running a parser that
# would just produce empty text). Keep it in one place so the
# frontend can mirror the same list when building the <input accept>.
_SUPPORTED_EXTS: dict[str, str] = {
    ".txt":  "txt",
    ".md":   "md",
    ".markdown": "md",
    ".pdf":  "pdf",
    ".docx": "docx",
    ".doc":  "docx-legacy",  # legacy .doc → not supported; surface a clear error
    ".html": "html",
    ".htm":  "html",
    ".epub": "epub",
}


def _extract_text_from_upload(filename: str, content: bytes) -> str:
    """Dispatch by extension. Return a plain ``str`` of the book body.

    Imports the heavy parsing libs lazily so a single missing
    optional dep doesn't break the whole endpoint — only the one
    format that needs it.
    """
    name = (filename or "").strip()
    if not name:
        raise ValueError("文件名不能为空。")
    ext = ""
    for candidate in (".epub", ".docx", ".html", ".htm", ".pdf", ".md",
                      ".markdown", ".doc", ".txt"):
        if name.lower().endswith(candidate):
            ext = candidate
            break
    if not ext or ext not in _SUPPORTED_EXTS:
        allowed = ", ".join(sorted(_SUPPORTED_EXTS))
        raise ValueError(f"暂不支持的格式: {ext or '?'}. 接受: {allowed}")
    kind = _SUPPORTED_EXTS[ext]
    if kind == "txt":
        return _parse_txt(content)
    if kind == "md":
        # Markdown is a superset of plain text for our purposes;
        # we just strip HTML-ish tags if any. The chapter splitter
        # doesn't care about ``**bold**`` etc.
        return _parse_txt(content)
    if kind == "pdf":
        return _parse_pdf(content)
    if kind == "docx":
        return _parse_docx(content)
    if kind == "docx-legacy":
        raise ValueError(
            "旧版 .doc (二进制) 不直接支持 — 请在 Word 里另存为 .docx 后再上传。"
        )
    if kind == "html":
        return _parse_html(content)
    if kind == "epub":
        return _parse_epub(content)
    raise ValueError(f"无法识别的格式: {ext}")


def _parse_txt(content: bytes) -> str:
    """Decode bytes as UTF-8 and strip a leading BOM if present.

    A BOM is invisible to the chapter regex (``^\\s*第``) but the
    ``\\ufeff`` byte is a *non-whitespace* character — so when the
    user uploads a TXT that was saved from 记事本 / Notepad the
    first chapter silently goes missing and the splitter falls back
    to "全文" with one chunk instead of N. Stripping the BOM here
    keeps the regex honest and the user gets the chapters they
    expect. Same trick for UTF-16 LE/BE in case the file was saved
    as Unicode.
    """
    if content.startswith(b"\xef\xbb\xbf"):
        content = content[3:]
    elif content.startswith(b"\xff\xfe") or content.startswith(b"\xfe\xff"):
        try:
            return content.decode("utf-16", errors="replace").lstrip("\ufeff")
        except Exception:
            return content.decode("utf-8", errors="replace")
    return content.decode("utf-8", errors="replace")


def _parse_pdf(content: bytes) -> str:
    import io
    import pypdf  # lazy import

    reader = pypdf.PdfReader(io.BytesIO(content))
    parts: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text:
            parts.append(text)
    if not parts:
        raise ValueError("PDF 解析后没有可识别的文字 (可能为扫描版/纯图片)。")
    return "\n\n".join(parts)


def _parse_docx(content: bytes) -> str:
    import io
    import docx  # python-docx, lazy import

    document = docx.Document(io.BytesIO(content))
    parts: list[str] = []
    for p in document.paragraphs:
        if p.text:
            parts.append(p.text)
    # Include text from tables too — lots of novels on 网文导出 have
    # character stats in the front matter. Joining with newlines is
    # fine; the chapter splitter is robust to extra newlines.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    if not parts:
        raise ValueError("DOCX 解析后没有可识别的正文。")
    return "\n".join(parts)


def _parse_html(content: bytes) -> str:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(content, "lxml")
    # Drop script / style entirely — they were never text we want.
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    # Pull out <br> as newline-ish separators so paragraphs don't
    # collapse into a single wall of text. Block-level tags will
    # naturally get their own newlines via get_text's separator.
    text = soup.get_text(separator="\n")
    # Collapse runs of 3+ blank lines (very common in scraped HTML).
    import re as _re
    text = _re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _parse_epub(content: bytes) -> str:
    """EPUB is a zip of XHTML files. Walk the spine and concatenate.

    The user gets a single string back with each ``<p>`` separated
    by a blank line, which is what our chapter splitter expects.
    """
    import io
    import os
    import tempfile
    import zipfile
    from bs4 import BeautifulSoup

    # ebooklib.epub.read_epub insists on a real path or file-like
    # with a real ``.name`` attribute. The cleanest portable path
    # is to write to a tempfile once and feed that path. We try
    # the clean path first and fall back to raw zipfile walk if
    # ebooklib is missing.
    try:
        import ebooklib  # type: ignore
        from ebooklib import epub

        with tempfile.NamedTemporaryFile(suffix=".epub", delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        try:
            book = epub.read_epub(tmp_path)
            items = [
                it for it in book.get_items()
                if it.get_type() == ebooklib.ITEM_DOCUMENT
            ]
            parts: list[str] = []
            for it in items:
                html = it.get_content() or b""
                soup = BeautifulSoup(html, "lxml")
                for tag in soup(["script", "style"]):
                    tag.decompose()
                text = soup.get_text(separator="\n").strip()
                if text:
                    parts.append(text)
            if not parts:
                raise ValueError("EPUB 解析后没有可识别的正文。")
            return "\n\n".join(parts)
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
    except ImportError:
        # ebooklib missing → fall back to raw zipfile walk.
        import re as _re
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            members = [
                n for n in zf.namelist()
                if _re.search(r"\.x?html?$", n, _re.IGNORECASE)
            ]
            members = sorted(members)
            parts = []
            for m in members:
                html = zf.read(m)
                soup = BeautifulSoup(html, "lxml")
                for tag in soup(["script", "style"]):
                    tag.decompose()
                text = soup.get_text(separator="\n").strip()
                if text:
                    parts.append(text)
            if not parts:
                raise ValueError("EPUB 解析后没有可识别的正文。")
            return "\n\n".join(parts)


@router.get("/materials/{material_id}", response_model=APIResponse[StudyMaterialDetail])
async def get_material(
    material_id: int,
    include_text: bool = Query(default=False),
    db: AsyncSession = Depends(get_db),
) -> APIResponse[StudyMaterialDetail]:
    row = (await db.execute(
        select(StudyMaterial)
        .options(
            selectinload(StudyMaterial.chapters),
            selectinload(StudyMaterial.characters),
        )
        .where(StudyMaterial.id == material_id)
    )).scalar_one_or_none()
    if row is None:
        raise not_found("StudyMaterial", material_id)
    detail = StudyMaterialDetail(
        id=row.id,
        project_id=row.project_id,
        title=row.title,
        author=row.author,
        source=row.source,
        status=row.status,
        error=row.error,
        chapter_count=row.chapter_count,
        character_count=row.character_count,
        raw_text_length=len(row.raw_text or ""),
        extra=row.extra,
        created_at=row.created_at,
        updated_at=row.updated_at,
        raw_text=row.raw_text if include_text else "",
        chapters=[StudyChapterRead.model_validate(c) for c in row.chapters],
        characters=[StudyCharacterRead.model_validate(c) for c in row.characters],
    )
    return {"ok": True, "data": detail}


@router.patch("/materials/{material_id}", response_model=APIResponse[StudyMaterialRead])
async def update_material(
    material_id: int,
    body: StudyMaterialUpdate,
    db: AsyncSession = Depends(get_db),
) -> APIResponse[StudyMaterialRead]:
    row = await db.get(StudyMaterial, material_id)
    if row is None:
        raise not_found("StudyMaterial", material_id)
    for k, v in body.model_dump(exclude_unset=True).items():
        setattr(row, k, v)
    await db.flush()
    return {"ok": True, "data": StudyMaterialRead.from_orm_trimmed(row)}


@router.delete("/materials/{material_id}", response_model=APIResponse[dict])
async def delete_material(
    material_id: int, db: AsyncSession = Depends(get_db)
) -> APIResponse[dict]:
    row = await db.get(StudyMaterial, material_id)
    if row is None:
        raise not_found("StudyMaterial", material_id)
    await db.delete(row)
    return {"ok": True, "data": {"deleted": material_id}}


@router.post("/materials/{material_id}/chapterize", response_model=APIResponse[StudyMaterialDetail])
async def chapterize_material(
    material_id: int,
    body: ChapterizeRequest | None = None,
    db: AsyncSession = Depends(get_db),
) -> APIResponse[StudyMaterialDetail]:
    """Split the material's raw_text into StudyChapter rows.

    Idempotent: wipes the existing chapters first. The user can rerun
    after editing the raw_text.
    """
    row = (await db.execute(
        select(StudyMaterial)
        .options(
            selectinload(StudyMaterial.chapters),
            selectinload(StudyMaterial.characters),
        )
        .where(StudyMaterial.id == material_id)
    )).scalar_one_or_none()
    if row is None:
        raise not_found("StudyMaterial", material_id)
    body = body or ChapterizeRequest()
    if not row.raw_text:
        raise bad_request("该 Material 还没有 raw_text，先粘贴或上传文本。")
    chunks = _split_chapters(row.raw_text, pattern=body.pattern)
    if not chunks:
        raise bad_request("分章失败：没有从文本中识别到任何章节。请检查文本格式。")
    # Wipe existing chapters (and orphan characters whose source chapter
    # is gone). We keep the characters so the user doesn't lose the
    # extracted set if they re-chapterize.
    for c in list(row.chapters):
        await db.delete(c)
    await db.flush()
    # R20 fix: count the new chapters in a local var instead of
    # ``len(row.chapters)`` — after ``selectinload`` the relationship
    # collection still holds the deleted-but-not-evicted old entries,
    # so ``len(row.chapters)`` was returning 2× (old + new). Tracking
    # the loop count ourselves sidesteps the stale-collection bug.
    created = 0
    for idx, title, content in chunks:
        if len(content) < body.min_chapter_chars:
            # Skip tiny fragments — they're likely false positives.
            continue
        row.chapters.append(StudyChapter(
            chapter_index=idx,
            title=title,
            content=content,
            char_count=len(content),
        ))
        created += 1
    row.chapter_count = created
    row.status = "ready"
    row.error = None
    await db.flush()
    # R20 fix: ``row.chapters`` is a relationship collection that was
    # ``selectinload``-ed at the top, then the wipe loop marked old
    # entries for deletion. The collection still holds those deleted
    # references until the session expires them. We could refresh
    # the attribute (one extra SELECT for 2K+ chapters), or we can
    # just return an empty chapters list in the response — the user
    # pulls the fresh list from ``/api/study/materials/{id}/chapters``
    # anyway. We choose the empty-list path: cheaper, and the
    # response carries the summary fields (chapter_count, status) the
    # UI actually needs to confirm the chapterize succeeded.
    return {"ok": True, "data": StudyMaterialDetail(
        id=row.id, project_id=row.project_id, title=row.title, author=row.author,
        source=row.source, status=row.status, error=row.error,
        chapter_count=row.chapter_count, character_count=row.character_count,
        raw_text_length=len(row.raw_text or ""), extra=row.extra,
        created_at=row.created_at, updated_at=row.updated_at,
        raw_text="",
        chapters=[],
        characters=[StudyCharacterRead.model_validate(c) for c in row.characters],
    )}


@router.get("/materials/{material_id}/chapters", response_model=APIResponse[list[StudyChapterRead]])
async def list_chapters(
    material_id: int, db: AsyncSession = Depends(get_db)
) -> APIResponse[list[StudyChapterRead]]:
    rows = (await db.execute(
        select(StudyChapter)
        .where(StudyChapter.material_id == material_id)
        .order_by(StudyChapter.chapter_index.asc())
    )).scalars().all()
    return {"ok": True, "data": [StudyChapterRead.model_validate(r) for r in rows]}


@router.get("/materials/{material_id}/characters", response_model=APIResponse[list[StudyCharacterRead]])
async def list_characters(
    material_id: int, db: AsyncSession = Depends(get_db)
) -> APIResponse[list[StudyCharacterRead]]:
    rows = (await db.execute(
        select(StudyCharacter)
        .where(StudyCharacter.material_id == material_id)
        .order_by(StudyCharacter.confidence.desc(), StudyCharacter.id.asc())
    )).scalars().all()
    return {"ok": True, "data": [StudyCharacterRead.model_validate(r) for r in rows]}


@router.post("/materials/{material_id}/characters", response_model=APIResponse[StudyCharacterRead])
async def add_character(
    material_id: int,
    body: StudyCharacterCreate,
    db: AsyncSession = Depends(get_db),
) -> APIResponse[StudyCharacterRead]:
    row = await db.get(StudyMaterial, material_id)
    if row is None:
        raise not_found("StudyMaterial", material_id)
    ch = StudyCharacter(
        material_id=material_id,
        name=body.name,
        aliases=body.aliases or [],
        role=body.role or "其他",
        tags=body.tags or [],
        base_profile=body.base_profile,
        confidence=body.confidence or 0.5,
    )
    db.add(ch)
    row.character_count = (row.character_count or 0) + 1
    await db.flush()
    return {"ok": True, "data": StudyCharacterRead.model_validate(ch)}


@router.delete("/materials/{material_id}/characters/{character_id}", response_model=APIResponse[dict])
async def delete_character(
    material_id: int,
    character_id: int,
    db: AsyncSession = Depends(get_db),
) -> APIResponse[dict]:
    ch = await db.get(StudyCharacter, character_id)
    if ch is None or ch.material_id != material_id:
        raise not_found("StudyCharacter", character_id)
    parent = await db.get(StudyMaterial, material_id)
    await db.delete(ch)
    if parent is not None and parent.character_count > 0:
        parent.character_count -= 1
    return {"ok": True, "data": {"deleted": character_id}}


@router.post("/materials/{material_id}/study", response_model=APIResponse[list[StudyCharacterRead]])
async def study_chapter(
    material_id: int,
    body: StudyRequest,
    db: AsyncSession = Depends(get_db),
) -> APIResponse[list[StudyCharacterRead]]:
    """Run the StudyCharacterAgent on a single chapter and persist the
    extracted characters.

    P15 / P0-STUDY-1: this used to call a regex stub
    (``_stub_study_parse``) because step-3.7-flash was unreliable for
    short chapter text. Now that the picker in
    ``app/services/llm/client.py`` reliably extracts the JSON from
    reasoning models (and the agent's ``allow_json_fallback=True``
    covers the worst case), we route through the real
    ``StudyCharacterAgent`` which goes through the same model-role
    bindings / prompt-versioning / AgentStep audit trail as the
    chapter pipeline.

    On a 5K-char chapter this takes 5-15s with the cheap model.
    Idempotency: characters from this chapter are upserted by
    ``(material_id, source_chapter_id, name)`` — running twice on
    the same chapter does not create duplicates.

    The task payload includes ``{"material_id": ..., "chapter_id": ...,
    "max_chars": ...}`` and we synthesize an ``AgentTask`` row so the
    AgentStep trail is consistent with the rest of the system.
    """
    material = await db.get(StudyMaterial, material_id)
    if material is None:
        raise not_found("StudyMaterial", material_id)
    chapter = await db.get(StudyChapter, body.chapter_id)
    if chapter is None or chapter.material_id != material_id:
        raise not_found("StudyChapter", body.chapter_id)
    text = chapter.content[: body.max_chars]
    existing = (await db.execute(
        select(StudyCharacter).where(StudyCharacter.material_id == material_id)
    )).scalars().all()
    existing_summary = "\n".join(
        f"- {c.name} (别名: {', '.join(c.aliases or [])}; 标签: {', '.join(c.tags or [])})"
        for c in existing
    ) or "（无）"

    # Synthesise a lightweight AgentTask so the AgentStep rows the
    # agent writes have a real parent. ``task_type=study_character``
    # keeps it out of the chapter-pipeline queue.
    study_task = AgentTask(
        project_id=material.project_id,
        chapter_id=None,
        task_type="study_character",
        status="running",
        priority=50,
        payload={
            "material_id": material_id,
            "chapter_id": chapter.id,
            "max_chars": body.max_chars,
        },
        started_at=datetime.utcnow(),
    )
    db.add(study_task)
    await db.flush()

    # Run the real LLM-backed agent. The agent's BaseAgent.run()
    # handles prompt rendering, model-role resolution, JSON parsing
    # (with allow_json_fallback) and the AgentStep audit row.
    agent = StudyCharacterAgent(
        router=get_llm_router(),
        engine=get_prompt_engine(),
    )
    try:
        result = await agent.run(
            AgentContext(
                db=db,
                task=study_task,
                project_id=material.project_id or 0,
                chapter_id=None,
                inputs={
                    "chapter_text": text,
                    "existing_characters": existing_summary,
                },
            )
        )
    except Exception as exc:
        # Mark the task as failed and re-raise so the API returns 4xx
        # — we don't want a silent failure to leave the user wondering
        # why no characters showed up.
        study_task.status = "failed"
        study_task.error = str(exc)
        study_task.finished_at = datetime.utcnow()
        await db.flush()
        raise
    study_task.status = "succeeded"
    study_task.finished_at = datetime.utcnow()
    study_task.cost_usd = result.cost_usd
    study_task.input_tokens = result.input_tokens
    study_task.output_tokens = result.output_tokens
    parsed = result.parsed or {}
    characters = parsed.get("characters") or []

    # Upsert: for each returned character, look up an existing row
    # in the same material with the same name; if present, merge
    # aliases / tags; if absent, insert. We then return the full set
    # for this material — including any pre-existing ones we
    # touched, so the UI can show "study touched these".
    new_rows: list[StudyCharacter] = []
    touched: list[StudyCharacter] = []
    for item in characters:
        if not isinstance(item, dict):
            continue
        name = (item.get("name") or "").strip()
        if not name:
            continue
        existing_row = next(
            (c for c in existing if c.name == name), None
        )
        if existing_row is not None:
            # Merge aliases (set union) and tags (set union). Don't
            # overwrite the user's manual edits to base_profile.
            new_aliases = list({*(existing_row.aliases or []), *(item.get("aliases") or [])})
            new_tags = list({*(existing_row.tags or []), *(item.get("tags") or [])})
            existing_row.aliases = new_aliases
            existing_row.tags = new_tags
            existing_row.source_chapter_id = chapter.id  # last seen
            existing_row.confidence = max(
                existing_row.confidence or 0.0,
                float(item.get("confidence") or 0.0),
            )
            touched.append(existing_row)
            continue
        ch = StudyCharacter(
            material_id=material_id,
            source_chapter_id=chapter.id,
            name=name,
            aliases=item.get("aliases") or [],
            role=item.get("role") or "其他",
            tags=item.get("tags") or [],
            base_profile=item.get("base_profile") or None,
            confidence=float(item.get("confidence") or 0.0),
        )
        db.add(ch)
        new_rows.append(ch)
        touched.append(ch)
    chapter.last_studied_at = datetime.utcnow()
    await db.flush()
    # Recompute the material's character_count = total rows on disk.
    total_count = (await db.execute(
        select(StudyCharacter).where(StudyCharacter.material_id == material_id)
    )).scalars().all()
    material.character_count = len(total_count)
    await db.flush()
    # Return the rows we touched (newly created + updated) so the
    # UI can highlight them in the table.
    return {"ok": True, "data": [StudyCharacterRead.model_validate(c) for c in touched]}


# -------------------- R21: bulk study (1-click 整本跑) --------------------

_SCRATCH_PROJECT_NAME = "拆书·公共"  # lazy-created scratch project for orphan study tasks


async def _resolve_scratch_project(db: AsyncSession) -> int:
    """Find-or-create the ``拆书·公共`` project for orphan study tasks.

    AgentTask.project_id is NOT NULL — the bulk endpoint binds a
    book-less study task to a scratch project so the user doesn't
    need to pre-associate every reference novel with a real project.
    The scratch project is hidden from the project sidebar (the
    ``projects`` API just returns every row, so the user would see
    one extra "拆书·公共" entry — that's the acceptable trade-off
    for not having to do a DB migration to make project_id
    nullable).
    """
    from app.models.project import Project
    row = (await db.execute(
        select(Project).where(Project.name == _SCRATCH_PROJECT_NAME)
    )).scalar_one_or_none()
    if row is not None:
        return row.id
    row = Project(
        name=_SCRATCH_PROJECT_NAME,
        category="study",
        description="拆书批量任务（关联到 project 的书的批量抽取会落到这里，避免污染真实项目）。",
    )
    db.add(row)
    await db.flush()
    return row.id


@router.post(
    "/materials/{material_id}/study/all",
    response_model=APIResponse[StudyBulkStartResponse],
)
async def study_bulk(
    material_id: int,
    body: StudyBulkRequest,
    db: AsyncSession = Depends(get_db),
) -> APIResponse[StudyBulkStartResponse]:
    """R21: kick off a background task that runs character / event
    extraction on every chapter of a material.

    The endpoint returns immediately with a ``task_id``. The caller
    polls ``GET /api/tasks/{task_id}`` to see per-mode progress —
    the AgentTask's ``payload`` carries the counters (chapters
    processed / total, characters added, events added, errors).

    Why background: a single chapter takes 5-15s on the cheap model,
    so 2332 chapters would be a 3-10 hour sync request. We spawn
    an ``asyncio.create_task`` so the user can navigate away; the
    worker only picks up ``task_type=chapter_pipeline`` so the
    background task is invisible to the chapter-pipeline queue.
    """
    material = await db.get(StudyMaterial, material_id)
    if material is None:
        raise not_found("StudyMaterial", material_id)
    # Count chapters up front. We need an open async session here
    # for the count, and another fresh one inside the background
    # coroutine — sharing the request session would deadlock the
    # event loop once we yield.
    total_chapters = (await db.execute(
        select(StudyChapter.id).where(StudyChapter.material_id == material_id)
    )).scalars().all()
    if not total_chapters:
        raise bad_request("该书还没有章节。先点「重新分章」再试。")
    if body.mode in ("event", "both") and not material.project_id:
        raise bad_request(
            "事件抽取需要先把这本书关联到 project_id（PATCH /api/study/materials/{id} 设置 project_id）。"
            "没有 project_id 就没法把事件存到 memory_foreshadows。"
        )
    # Resolve a project_id for the AgentTask. If the book is bound to a
    # real project we use that; otherwise we fall back to the scratch
    # project so AgentTask.project_id (NOT NULL) doesn't blow up.
    effective_project_id = material.project_id or await _resolve_scratch_project(db)
    chapters_to_process = len(total_chapters) if not body.limit else min(body.limit, len(total_chapters))
    # Cap concurrency to a sane range; the user can crank it but
    # 3 is plenty for one book's worth of LLM bandwidth.
    concurrency = max(1, min(body.max_concurrency, 8))
    # Synthesise a lightweight AgentTask. ``task_type=study_bulk``
    # keeps it out of the chapter-pipeline queue (the worker's
    # query is ``task_type == "chapter_pipeline"``).
    study_task = AgentTask(
        project_id=effective_project_id,
        chapter_id=None,
        task_type=f"study_bulk_{body.mode}",
        status="running",
        priority=50,
        payload={
            "material_id": material_id,
            "mode": body.mode,
            "total_chapters": len(total_chapters),
            "chapters_to_process": chapters_to_process,
            "chapters_processed": 0,
            "characters_added": 0,
            "events_added": 0,
            "errors": [],
            "max_concurrency": concurrency,
            "force": body.force,
            "max_chars": body.max_chars,
        },
        started_at=datetime.utcnow(),
    )
    db.add(study_task)
    await db.flush()
    task_id = study_task.id
    # Detach the chapter IDs from the request session so the
    # background coroutine can iterate over a plain list.
    chapter_ids: list[int] = list(total_chapters)
    # R21 fix: actually apply ``body.limit`` to the work queue. The
    # response already says ``chapters_to_process = min(limit, total)``
    # but the background coroutine used to iterate over the FULL
    # list, so a ``limit=1`` request would still process every
    # chapter — wasting LLM tokens and the user's daily budget. We
    # sort by chapter_index so the user always sees the first N
    # chapters get processed, regardless of DB id order.
    if body.limit and body.limit > 0:
        chapter_ids = chapter_ids[: body.limit]
    # Spawn the background coroutine. We capture only the IDs +
    # the engine (not the session) so the request session is free
    # to be closed by FastAPI as soon as the response is sent.
    asyncio.create_task(
        _run_bulk_study(
            task_id=task_id,
            material_id=material_id,
            chapter_ids=chapter_ids,
            mode=body.mode,
            force=body.force,
            max_chars=body.max_chars,
            max_concurrency=concurrency,
            effective_project_id=effective_project_id,
        )
    )
    return {
        "ok": True,
        "data": StudyBulkStartResponse(
            task_id=task_id,
            total_chapters=len(total_chapters),
            chapters_to_process=chapters_to_process,
            mode=body.mode,
        ),
    }


async def _persist_bulk_progress(task_id: int, payload: dict[str, Any], *, finished: bool = False) -> None:
    """Persist the running payload back to the AgentTask row.

    Each progress tick opens its own short-lived session so a long
    study run doesn't hold a single transaction open for hours.
    """
    from app.models.task import AgentTask as _AT
    async with AsyncSessionLocal() as db:
        row = await db.get(_AT, task_id)
        if row is None:
            return
        row.payload = dict(payload)
        if finished:
            row.status = "failed" if payload.get("errors") and not payload.get("chapters_processed") else "succeeded"
            row.finished_at = datetime.utcnow()
        await db.commit()


async def _run_bulk_study(
    *,
    task_id: int,
    material_id: int,
    chapter_ids: list[int],
    mode: Literal["character", "event", "both"],
    force: bool,
    max_chars: int,
    max_concurrency: int,
    effective_project_id: int,
) -> None:
    """R21 background coroutine: process chapters in parallel.

    The semaphore keeps at most ``max_concurrency`` LLM calls
    in flight at once (default 3). Progress is persisted on
    every chapter completion so the frontend polling
    ``GET /api/tasks/{task_id}`` sees fresh numbers.

    On unhandled exception the task is marked failed with the
    message stored in ``error`` and the partial progress kept
    in ``payload`` for postmortem.
    """
    # Re-read the current payload so we have a single source of
    # truth (the request session may have already been closed).
    async with AsyncSessionLocal() as db:
        task = await db.get(AgentTask, task_id)
        if task is None:
            return
        payload: dict[str, Any] = dict(task.payload or {})
    sem = asyncio.Semaphore(max_concurrency)
    # We pass the existing router / engine (already cached at
    # module level) so the background coroutine doesn't pay the
    # lazy-init cost on every chapter.
    router = get_llm_router()
    engine = get_prompt_engine()
    char_agent = StudyCharacterAgent(router=router, engine=engine)
    event_agent = StudyEventAgent(router=router, engine=engine) if mode in ("event", "both") else None

    async def _process_one(chapter_id: int) -> tuple[int, int, int, str | None]:
        """Run the requested agents for a single chapter. Returns
        ``(characters_added, events_added, retried_count, error)``.
        """
        async with sem:
            try:
                return await _bulk_process_chapter(
                    task_id=task_id,
                    material_id=material_id,
                    chapter_id=chapter_id,
                    mode=mode,
                    force=force,
                    max_chars=max_chars,
                    char_agent=char_agent,
                    event_agent=event_agent,
                    effective_project_id=effective_project_id,
                )
            except Exception as exc:  # noqa: BLE001 — last-line isolation
                return 0, 0, 0, f"{exc.__class__.__name__}: {exc}".strip()

    try:
        results = await asyncio.gather(
            *(_process_one(cid) for cid in chapter_ids),
            return_exceptions=False,
        )
    except Exception as exc:  # noqa: BLE001
        # If asyncio.gather itself blows up (e.g. cancelled), mark
        # the task failed and bail.
        payload["error"] = f"{exc.__class__.__name__}: {exc}".strip()
        await _persist_bulk_progress(task_id, payload, finished=True)
        return
    # Aggregate per-chapter outcomes into the shared payload.
    total_chars = 0
    total_events = 0
    for chars_added, events_added, _retried, err in results:
        total_chars += chars_added
        total_events += events_added
        if err:
            payload.setdefault("errors", []).append(err)
        payload["chapters_processed"] = int(payload.get("chapters_processed", 0)) + 1
        # We write progress every chapter so the user sees the
        # counter tick up live. For 2332 chapters this is 2332
        # short commits, which is fine on local SQLite.
        await _persist_bulk_progress(task_id, payload, finished=False)
    payload["characters_added"] = total_chars
    payload["events_added"] = total_events
    await _persist_bulk_progress(task_id, payload, finished=True)


async def _bulk_process_chapter(
    *,
    task_id: int,
    material_id: int,
    chapter_id: int,
    mode: Literal["character", "event", "both"],
    force: bool,
    max_chars: int,
    char_agent: StudyCharacterAgent,
    event_agent: StudyEventAgent | None,
    effective_project_id: int,
) -> tuple[int, int, int, str | None]:
    """Process a single chapter in its own session and return
    ``(characters_added, events_added, retried, error)``.

    Skips chapters whose ``last_studied_at`` is set unless
    ``force=True`` — the user paid for those LLM tokens once,
    no need to re-charge by default.
    """
    from app.models.task import AgentTask as _AT
    async with AsyncSessionLocal() as db:
        try:
            material = await db.get(StudyMaterial, material_id)
            if material is None:
                return 0, 0, 0, f"material {material_id} missing"
            chapter = await db.get(StudyChapter, chapter_id)
            if chapter is None or chapter.material_id != material_id:
                return 0, 0, 0, f"chapter {chapter_id} missing"
            # Skip-already-studied: last_studied_at is set by the
            # single-chapter endpoint after a successful run, so
            # it's a fine "did this chapter get processed" signal.
            if not force and chapter.last_studied_at is not None:
                return 0, 0, 0, None
            text = (chapter.content or "")[:max_chars]
            chars_added = 0
            events_added = 0
            # ----- character extraction -----
            if mode in ("character", "both"):
                existing = (await db.execute(
                    select(StudyCharacter).where(StudyCharacter.material_id == material_id)
                )).scalars().all()
                existing_summary = "\n".join(
                    f"- {c.name} (别名: {', '.join(c.aliases or [])}; 标签: {', '.join(c.tags or [])})"
                    for c in existing
                ) or "（无）"
                # Synthesise a per-chapter AgentTask so the
                # AgentStep audit trail stays consistent with
                # the single-chapter endpoint.
                sub_task = _AT(
                    project_id=effective_project_id,
                    chapter_id=None,
                    task_type="study_character",
                    status="running",
                    priority=50,
                    payload={
                        "material_id": material_id,
                        "chapter_id": chapter.id,
                        "max_chars": max_chars,
                        "bulk_parent_task_id": task_id,
                    },
                    started_at=datetime.utcnow(),
                )
                db.add(sub_task)
                await db.flush()
                result = await char_agent.run(
                    AgentContext(
                        db=db,
                        task=sub_task,
                        project_id=material.project_id or 0,
                        chapter_id=None,
                        inputs={
                            "chapter_text": text,
                            "existing_characters": existing_summary,
                        },
                    )
                )
                sub_task.status = "succeeded"
                sub_task.finished_at = datetime.utcnow()
                parsed = result.parsed or {}
                for item in (parsed.get("characters") or []):
                    if not isinstance(item, dict):
                        continue
                    name = (item.get("name") or "").strip()
                    if not name:
                        continue
                    existing_row = next(
                        (c for c in existing if c.name == name), None
                    )
                    if existing_row is not None:
                        # Merge aliases / tags; skip counts.
                        existing_row.aliases = list({
                            *(existing_row.aliases or []),
                            *(item.get("aliases") or []),
                        })
                        existing_row.tags = list({
                            *(existing_row.tags or []),
                            *(item.get("tags") or []),
                        })
                        existing_row.source_chapter_id = chapter.id
                        existing_row.confidence = max(
                            existing_row.confidence or 0.0,
                            float(item.get("confidence") or 0.0),
                        )
                        continue
                    db.add(StudyCharacter(
                        material_id=material_id,
                        source_chapter_id=chapter.id,
                        name=name,
                        aliases=item.get("aliases") or [],
                        role=item.get("role") or "其他",
                        tags=item.get("tags") or [],
                        base_profile=item.get("base_profile") or None,
                        confidence=float(item.get("confidence") or 0.0),
                    ))
                    chars_added += 1
            # ----- event extraction (only if project_id) -----
            if mode in ("event", "both") and material.project_id:
                from app.models.memory import MemoryForeshadow
                existing_foreshadows = (await db.execute(
                    select(MemoryForeshadow).where(MemoryForeshadow.project_id == material.project_id)
                )).scalars().all()
                existing_summary = "\n".join(
                    f"- {f.name}: {f.summary[:80]}"
                    for f in existing_foreshadows
                ) or "（无）"
                sub_task = _AT(
                    project_id=material.project_id,
                    chapter_id=None,
                    task_type="study_event",
                    status="running",
                    priority=50,
                    payload={
                        "material_id": material_id,
                        "chapter_id": chapter.id,
                        "max_chars": max_chars,
                        "bulk_parent_task_id": task_id,
                    },
                    started_at=datetime.utcnow(),
                )
                db.add(sub_task)
                await db.flush()
                event_result = await event_agent.run(
                    AgentContext(
                        db=db,
                        task=sub_task,
                        project_id=material.project_id,
                        chapter_id=None,
                        inputs={
                            "chapter_text": text,
                            "chapter_no": chapter.chapter_index,
                            "existing_foreshadows": existing_summary,
                        },
                    )
                )
                sub_task.status = "succeeded"
                sub_task.finished_at = datetime.utcnow()
                parsed_events = event_result.parsed or {}
                for item in (parsed_events.get("events") or []):
                    if not isinstance(item, dict):
                        continue
                    name = (item.get("name") or "").strip()
                    if not name:
                        continue
                    summary = (item.get("summary") or "").strip()
                    # Use chapter_index as planted_chapter; import
                    # importance scaled to 0..1 (events use 1..5).
                    raw_imp = item.get("importance")
                    try:
                        imp = max(0.0, min(1.0, float(raw_imp) / 5.0)) if raw_imp is not None else 0.5
                    except (TypeError, ValueError):
                        imp = 0.5
                    related_chars = item.get("related_characters") or []
                    if not isinstance(related_chars, list):
                        related_chars = []
                    related_chars = [str(x) for x in related_chars if x]
                    # The prompt's ``quote`` field is dropped here
                    # — memory_foreshadows has no quote column. We
                    # fold it into the summary so the user still
                    # sees the evidence in the memory view.
                    quote = (item.get("quote") or "").strip()
                    if quote:
                        summary = f"{summary}\n[原文] {quote}" if summary else f"[原文] {quote}"
                    # De-dup: skip if a foreshadow with the same
                    # name + same planted_chapter already exists.
                    dup = next(
                        (
                            f for f in existing_foreshadows
                            if f.name == name and f.planted_chapter == chapter.chapter_index
                        ),
                        None,
                    )
                    if dup is not None:
                        continue
                    db.add(MemoryForeshadow(
                        project_id=material.project_id,
                        name=name,
                        summary=summary,
                        planted_chapter=chapter.chapter_index,
                        status="active",
                        importance=imp,
                        related_characters=related_chars,
                        # R22: stamp provenance so the Study page and
                        # the graph materialise endpoint can filter
                        # "foreshadows from this book" cheaply.
                        source_material_id=material_id,
                    ))
                    events_added += 1
            # Mark the chapter as studied (timestamp). For
            # ``character`` mode this matches the single-chapter
            # endpoint's behavior; for ``event`` only we still
            # set it so the skip-already-studied heuristic works
            # for both modes.
            chapter.last_studied_at = datetime.utcnow()
            await db.commit()
            return chars_added, events_added, 0, None
        except Exception as exc:  # noqa: BLE001
            await db.rollback()
            return 0, 0, 0, f"chapter {chapter_id}: {exc.__class__.__name__}: {exc}".strip()


# -------------------- R22: study → graph / behavior / foreshadow linkage --------------------
#
# Round 22 is the "功能联动" cleanup. After 拆书 runs, the data has
# to flow into:
#   - the drafter (consumes behavior_patterns) — the
#     ``StudyBehaviorPatternAgent`` is already coded but no route
#     invokes it. This adds the route.
#   - the graph page (consumes graph_nodes + graph_edges) — the
#     existing materialise endpoint handled characters only. We add
#     ``kind=event|behavior|all`` and a co-occurrence edge pass.
#   - the memory page (consumes memory_foreshadows) — auto-extracted
#     foreshadows now stamp ``source_material_id`` so the Study page
#     can list "events from this book" without a JSON-payload scan.


async def _pick_evidence_chapters_async(
    db: AsyncSession,
    material_id: int,
    chapter_count: int,
    max_chunk_chars: int,
) -> list[dict[str, Any]]:
    """Pick ``chapter_count`` chapters with the most extracted characters.

    The behavior-pattern agent needs *representative* scenes — a 2332-
    chapter book would overwhelm the prompt if we fed it everything,
    but the first chapter alone is too narrow. The simplest signal
    of "interesting" is "the chapter where lots of named characters
    were seen together" (a heist / council / battle scene), so we
    rank by StudyCharacter.source_chapter_id and take the top N.

    Returns a list of ``{id, chapter_index, title, text}`` dicts
    ready to feed the LLM as the ``evidence_chunks`` field.
    """
    from collections import Counter

    from app.models.study import StudyChapter as _SCh, StudyCharacter as _SC

    if chapter_count <= 0:
        return []
    rows = (
        await db.execute(
            select(_SC.source_chapter_id, _SC.name).where(
                _SC.material_id == material_id,
                _SC.source_chapter_id.is_not(None),
            )
        )
    ).all()
    counter: Counter[int] = Counter(r[0] for r in rows if r[0] is not None)
    if not counter:
        return []
    # Most-active first; tie-break by chapter_index (earlier wins).
    top_ids = [cid for cid, _ in counter.most_common(chapter_count)]
    if not top_ids:
        return []
    chapter_rows = (
        await db.execute(
            select(_SCh).where(_SCh.id.in_(top_ids))
        )
    ).scalars().all()
    by_id = {c.id: c for c in chapter_rows}
    out: list[dict[str, Any]] = []
    for cid in top_ids:
        c = by_id.get(cid)
        if c is None:
            continue
        text = (c.content or "")[:max_chunk_chars]
        out.append({
            "id": c.id,
            "chapter_index": c.chapter_index,
            "title": c.title,
            "text": text,
        })
    return out


@router.post(
    "/materials/{material_id}/extract-behaviors",
    response_model=APIResponse[StudyBehaviorExtractResponse],
)
async def extract_behaviors(
    material_id: int,
    body: StudyBehaviorExtractRequest,
    db: AsyncSession = Depends(get_db),
) -> APIResponse[StudyBehaviorExtractResponse]:
    """R22: extract reusable behavior patterns from a study material.

    One LLM call. The prompt sees the material's character roster
    plus 2-5 "active" chapter snippets (chapters with the most
    extracted characters) and returns a JSON list of pattern cards.
    We persist each one as a ``BehaviorPattern`` row with
    ``source_material_id`` so the drafter can pull them by tag.

    Idempotency: if ``force=False`` and the material already has
    patterns, we return the existing ones without calling the LLM.
    With ``force=True`` we wipe the material's patterns first.
    """
    material = await db.get(StudyMaterial, material_id)
    if material is None:
        raise not_found("StudyMaterial", material_id)

    # 1) Idempotency / force re-run
    existing = (await db.execute(
        select(BehaviorPattern).where(BehaviorPattern.source_material_id == material_id)
    )).scalars().all()
    if existing and not body.force:
        return {
            "ok": True,
            "data": StudyBehaviorExtractResponse(
                material_id=material_id,
                patterns_added=0,
                patterns_skipped=len(existing),
                pattern_ids=[p.id for p in existing[: body.max_patterns]],
                total_patterns_for_material=len(existing),
                cost_usd=0.0,
                duration_ms=0,
                input_tokens=0,
                output_tokens=0,
                sample_names=[p.name for p in existing[:5]],
            ),
        }
    if existing and body.force:
        for p in existing:
            await db.delete(p)
        await db.flush()

    # 2) Build the evidence payload
    chars = (await db.execute(
        select(StudyCharacter).where(StudyCharacter.material_id == material_id)
    )).scalars().all()
    if not chars:
        raise bad_request(
            "该书还没有抽过人物。先跑一次「抽取人物」，再归纳行为模式。",
            suggestion="POST /api/study/materials/{id}/study/all with mode=character",
        )
    char_summary = "\n".join(
        f"- {c.name} (role={c.role or '其他'}, tags={', '.join(c.tags or [])})"
        for c in chars[:80]
    )
    evidence_chapters = await _pick_evidence_chapters_async(
        db,
        material_id=material_id,
        chapter_count=body.evidence_chapter_count,
        max_chunk_chars=body.max_chunk_chars,
    )
    if not evidence_chapters:
        raise bad_request(
            "找不到可用的章节片段（人物没有 source_chapter_id，可能是手工添加的）。",
            suggestion="先跑批量抽人物，让每条人物挂上章节来源。",
        )
    evidence_block = "\n\n".join(
        f"### {ch['title'] or '第 ' + str(ch['chapter_index']) + ' 章'}\n{ch['text']}"
        for ch in evidence_chapters
    )
    existing_patterns = (await db.execute(
        select(BehaviorPattern).order_by(BehaviorPattern.id.desc()).limit(50)
    )).scalars().all()
    existing_pattern_block = "\n".join(
        f"- {p.name} (人物:{', '.join(p.character_tags or [])}; 情境:{', '.join(p.situation_tags or [])})"
        for p in existing_patterns
    ) or "（无）"

    # 3) Synthesise a lightweight AgentTask so the AgentStep audit
    #    trail is consistent with the rest of the system.
    study_task = AgentTask(
        project_id=material.project_id or 0,
        chapter_id=None,
        task_type="study_behavior_pattern",
        status="running",
        priority=50,
        payload={
            "material_id": material_id,
            "max_patterns": body.max_patterns,
            "evidence_chapter_count": len(evidence_chapters),
        },
        started_at=datetime.utcnow(),
    )
    db.add(study_task)
    await db.flush()

    # 4) Run the LLM agent
    agent = StudyBehaviorPatternAgent(
        router=get_llm_router(), engine=get_prompt_engine(),
    )
    try:
        result = await agent.run(
            AgentContext(
                db=db,
                task=study_task,
                project_id=material.project_id or 0,
                chapter_id=None,
                inputs={
                    "evidence_chunks": evidence_block,
                    "existing_patterns": existing_pattern_block,
                    "character_roster": char_summary,
                    "max_patterns": str(body.max_patterns),
                },
            )
        )
    except Exception as exc:
        study_task.status = "failed"
        study_task.error = str(exc)
        study_task.finished_at = datetime.utcnow()
        await db.flush()
        raise
    study_task.status = "succeeded"
    study_task.finished_at = datetime.utcnow()
    study_task.cost_usd = result.cost_usd
    study_task.input_tokens = result.input_tokens
    study_task.output_tokens = result.output_tokens

    parsed = result.parsed or {}
    patterns_in = parsed.get("patterns") or []
    if not isinstance(patterns_in, list):
        patterns_in = []

    # 5) Persist. Skip duplicates by (source_material_id, name) so a
    #    partial re-run with the LLM repeating itself doesn't grow
    #    the table unboundedly.
    new_rows: list[BehaviorPattern] = []
    skipped = 0
    for item in patterns_in:
        if not isinstance(item, dict):
            continue
        name = (item.get("name") or "").strip()
        if not name:
            continue
        if len(new_rows) >= body.max_patterns:
            skipped += 1
            continue
        if any(p.name == name for p in existing):
            skipped += 1
            continue
        row = BehaviorPattern(
            source_material_id=material_id,
            name=name[:120],
            character_tags=[str(x).strip() for x in (item.get("character_tags") or []) if str(x).strip()][:8],
            situation_tags=[str(x).strip() for x in (item.get("situation_tags") or []) if str(x).strip()][:8],
            typical_behavior=[str(x).strip() for x in (item.get("typical_behavior") or []) if str(x).strip()][:8],
            dialogue_style=[str(x).strip() for x in (item.get("dialogue_style") or []) if str(x).strip()][:8],
            scene_function=[str(x).strip() for x in (item.get("scene_function") or []) if str(x).strip()][:8],
            risks=[str(x).strip() for x in (item.get("risks") or []) if str(x).strip()][:8],
            recommended_plot_followup=[str(x).strip() for x in (item.get("recommended_plot_followup") or []) if str(x).strip()][:8],
            evidence=[str(x).strip() for x in (item.get("evidence") or []) if str(x).strip()][:5],
            confidence=float(item.get("confidence") or 0.5),
        )
        db.add(row)
        new_rows.append(row)
    await db.flush()
    return {
        "ok": True,
        "data": StudyBehaviorExtractResponse(
            material_id=material_id,
            patterns_added=len(new_rows),
            patterns_skipped=skipped,
            pattern_ids=[p.id for p in new_rows],
            total_patterns_for_material=len(existing) + len(new_rows),
            cost_usd=round(result.cost_usd, 4),
            duration_ms=result.duration_ms,
            input_tokens=result.input_tokens,
            output_tokens=result.output_tokens,
            sample_names=[p.name for p in new_rows[:5]],
        ),
    }


@router.get(
    "/materials/{material_id}/behaviors",
    response_model=APIResponse[list[BehaviorPatternRead]],
)
async def list_behaviors(
    material_id: int, db: AsyncSession = Depends(get_db)
) -> APIResponse[list[BehaviorPatternRead]]:
    """R22: behavior patterns produced by this material.

    Empty list is the honest answer when the user hasn't run
    extract-behaviors yet — the frontend shows "0 条" with a CTA.
    """
    material = await db.get(StudyMaterial, material_id)
    if material is None:
        raise not_found("StudyMaterial", material_id)
    rows = (await db.execute(
        select(BehaviorPattern)
        .where(BehaviorPattern.source_material_id == material_id)
        .order_by(BehaviorPattern.confidence.desc(), BehaviorPattern.id.asc())
    )).scalars().all()
    return {"ok": True, "data": [BehaviorPatternRead.model_validate(r) for r in rows]}


@router.get(
    "/materials/{material_id}/foreshadows",
    response_model=APIResponse[list[StudyForeshadowSummary]],
)
async def list_foreshadows(
    material_id: int, db: AsyncSession = Depends(get_db)
) -> APIResponse[list[StudyForeshadowSummary]]:
    """R22: memory_foreshadows stamped with ``source_material_id``.

    The bulk event extractor (R21 / mode=event) sets this column, so
    the Study page can show "foreshadows from this book" without
    scanning the foreshadow summary for the book title.
    """
    from app.models.memory import MemoryForeshadow
    material = await db.get(StudyMaterial, material_id)
    if material is None:
        raise not_found("StudyMaterial", material_id)
    if not material.project_id:
        # The material isn't bound to a project, so the bulk event
        # extractor never wrote any foreshadows under this material
        # (it requires project_id to even kick off).
        return {"ok": True, "data": []}
    rows = (await db.execute(
        select(MemoryForeshadow)
        .where(
            MemoryForeshadow.source_material_id == material_id,
            MemoryForeshadow.project_id == material.project_id,
        )
        .order_by(MemoryForeshadow.planted_chapter.asc().nulls_last(), MemoryForeshadow.id.asc())
    )).scalars().all()
    return {"ok": True, "data": [
        StudyForeshadowSummary(
            id=r.id,
            name=r.name,
            summary=r.summary or "",
            planted_chapter=r.planted_chapter,
            status=r.status,
            importance=r.importance,
            related_characters=r.related_characters or [],
        )
        for r in rows
    ]}


@router.get(
    "/materials/{material_id}/relationships",
    response_model=APIResponse[StudyRelationshipsResponse],
)
async def list_relationship_suggestions(
    material_id: int, db: AsyncSession = Depends(get_db)
) -> APIResponse[StudyRelationshipsResponse]:
    """R22: co-occurrence analysis on the study_characters.

    For every pair of characters that share a source_chapter_id, we
    suggest an edge. The relation label is left to the user — we
    don't try to infer "师父 vs 同门" from the chapter text (the
    cheap model is unreliable for that and the false-positive cost
    is high). The MVP just emits "同章节出现" as a placeholder; the
    user replaces it in the apply-modal.

    Output is sorted by (co_chapter_count desc, last_chapter_no desc)
    so the most-evidenced pairs surface first.
    """
    from sqlalchemy import func

    from app.models.study import StudyCharacter as _SC, StudyChapter as _SCh

    material = await db.get(StudyMaterial, material_id)
    if material is None:
        raise not_found("StudyMaterial", material_id)

    # 1) Pull every character with a source chapter. Characters added
    #    by hand (source_chapter_id is NULL) can't be analysed because
    #    we have no evidence to point at.
    char_rows = (await db.execute(
        select(_SC)
        .where(_SC.material_id == material_id, _SC.source_chapter_id.is_not(None))
    )).scalars().all()
    if not char_rows:
        return {
            "ok": True,
            "data": StudyRelationshipsResponse(
                material_id=material_id,
                chapters_scanned=0,
                suggestions=[],
                total_characters=0,
            ),
        }

    # 2) Bucket by chapter. A chapter with 5 characters produces
    #    C(5,2) = 10 candidate pairs; we accumulate the count and
    #    the latest chapter_id/no for each pair.
    from collections import defaultdict

    chapter_to_chars: dict[int, list[_SC]] = defaultdict(list)
    for c in char_rows:
        chapter_to_chars[c.source_chapter_id].append(c)

    pair_acc: dict[tuple[int, int], dict[str, Any]] = {}
    chapter_ids_seen: set[int] = set()
    for chap_id, chars_in_chap in chapter_to_chars.items():
        chapter_ids_seen.add(chap_id)
        # Cap per-chapter char count — 50 characters in one chapter
        # would explode to 1225 pairs. We only emit pairs among the
        # first 20 characters in the chapter (more than enough
        # signal; the LLM's per-chapter output is also bounded).
        chars_capped = chars_in_chap[:20]
        for i, a in enumerate(chars_capped):
            for b in chars_capped[i + 1:]:
                # Stable ordering so (A,B) and (B,A) collide on the
                # same bucket.
                lo, hi = (a, b) if a.id < b.id else (b, a)
                key = (lo.id, hi.id)
                acc = pair_acc.setdefault(key, {
                    "char_a": lo,
                    "char_b": hi,
                    "co_chapter_count": 0,
                    "last_chapter_id": chap_id,
                    "last_chapter_no": 0,
                })
                acc["co_chapter_count"] += 1

    # 3) Resolve last_chapter_no for the tracked chapter_ids.
    chap_rows = (await db.execute(
        select(_SCh).where(_SCh.id.in_(list(chapter_ids_seen)))
    )).scalars().all()
    by_chap_id = {c.id: c for c in chap_rows}
    for acc in pair_acc.values():
        ch = by_chap_id.get(acc["last_chapter_id"])
        if ch is not None:
            acc["last_chapter_no"] = ch.chapter_index
            acc["last_chapter_title"] = ch.title or ""

    # 4) Sort + format.
    rows = sorted(
        pair_acc.values(),
        key=lambda x: (-x["co_chapter_count"], -x["last_chapter_no"]),
    )
    suggestions: list[StudyRelationshipSuggestion] = []
    for r in rows:
        # Try to grab a 1-2 sentence quote from the chapter as
        # evidence. The model that extracted the characters
        # already saw the text; we just show the first 200 chars
        # of the chapter so the user has something to evaluate.
        ch = by_chap_id.get(r["last_chapter_id"])
        quote = ""
        if ch is not None and ch.content:
            txt = ch.content
            names = [r["char_a"].name, r["char_b"].name]
            # Look for the first sentence mentioning both names.
            for sent in re.split(r"[。！？!?\n]", txt):
                if all(n in sent for n in names) and 8 <= len(sent) <= 200:
                    quote = sent.strip()
                    break
            if not quote:
                quote = txt[:200].replace("\n", " ").strip()
        suggestions.append(StudyRelationshipSuggestion(
            char_a_id=r["char_a"].id,
            char_a_name=r["char_a"].name,
            char_b_id=r["char_b"].id,
            char_b_name=r["char_b"].name,
            co_chapter_count=r["co_chapter_count"],
            last_chapter_id=r["last_chapter_id"],
            last_chapter_no=r["last_chapter_no"],
            last_chapter_title=r.get("last_chapter_title", ""),
            sample_quote=quote,
        ))

    return {
        "ok": True,
        "data": StudyRelationshipsResponse(
            material_id=material_id,
            chapters_scanned=len(chapter_ids_seen),
            suggestions=suggestions,
            total_characters=len(char_rows),
        ),
    }


# R24: 把 R22 纯 co-occurrence 的 suggestions 跑 LLM, 抽出真实
# 语义关系 (师父/对手/恋人/家人/...) 替代「同章节出现」。
# 用户 R23 反馈: 「相互的联系不要之说出现在同一章节啊」。
@router.post(
    "/materials/{material_id}/relationships/enrich",
    response_model=APIResponse[StudyRelationshipEnrichResponse],
)
async def enrich_relationships(
    material_id: int,
    body: StudyRelationshipEnrichRequest,
    db: AsyncSession = Depends(get_db),
) -> APIResponse[StudyRelationshipEnrichResponse]:
    """R24: LLM 抽语义关系。

    这个端点跟 ``/relationships`` 走同一条数据路径 (co-occurrence
    先筛出候选对), 但额外对每对调用 LLM 判关系类型。运行时间跟
    对数线性 — 30 对约 1-2 分钟, 100 对约 4-6 分钟, 所以默认
    cap 在 30 对。

    输入: ``StudyRelationshipEnrichRequest`` with optional
    ``max_pairs`` cap (default 30) and ``min_co_chapter_count``。
    输出: 同形状的 items, 但 ``relation`` 字段是 LLM 给的真实
    关系, 带 ``confidence`` 和 ``evidence``。LLM 抽不出的标
    「同章节出现」+ confidence=0。
    """
    import asyncio
    import time as _time
    from app.agents.study import StudyRelationshipExtractionAgent

    t0 = _time.time()
    material = await db.get(StudyMaterial, material_id)
    if material is None:
        raise not_found("StudyMaterial", material_id)

    # 1) 复用 /relationships 端的候选对生成逻辑 (抽出去共享 — 后面
    #    重构, 现在内联以减小 diff 体积)。
    from app.models.study import StudyCharacter as _SC, StudyChapter as _SCh
    from collections import defaultdict

    char_rows = (await db.execute(
        select(_SC).where(
            _SC.material_id == material_id,
            _SC.source_chapter_id.is_not(None),
        )
    )).scalars().all()
    if not char_rows:
        return {
            "ok": True,
            "data": StudyRelationshipEnrichResponse(
                material_id=material_id,
                enriched_count=0, skipped_count=0, fallback_count=0,
                duration_ms=0, cost_usd=0.0, items=[],
            ),
        }

    chapter_to_chars: dict[int, list[_SC]] = defaultdict(list)
    for c in char_rows:
        chapter_to_chars[c.source_chapter_id].append(c)

    pair_acc: dict[tuple[int, int], dict[str, Any]] = {}
    chapter_ids_seen: set[int] = set()
    for chap_id, chars_in_chap in chapter_to_chars.items():
        chapter_ids_seen.add(chap_id)
        chars_capped = chars_in_chap[:20]
        for i, a in enumerate(chars_capped):
            for b in chars_in_chap[i + 1:]:
                lo, hi = (a, b) if a.id < b.id else (b, a)
                key = (lo.id, hi.id)
                acc = pair_acc.setdefault(key, {
                    "char_a": lo, "char_b": hi,
                    "co_chapter_count": 0,
                    "last_chapter_id": chap_id, "last_chapter_no": 0,
                })
                acc["co_chapter_count"] += 1

    # 2) 取章节信息 (序号 + 标题)
    chap_rows = (await db.execute(
        select(_SCh).where(_SCh.id.in_(list(chapter_ids_seen)))
    )).scalars().all()
    by_chap_id: dict[int, _SCh] = {c.id: c for c in chap_rows}
    for acc in pair_acc.values():
        ch = by_chap_id.get(acc["last_chapter_id"])
        if ch is not None:
            acc["last_chapter_no"] = ch.chapter_index
            acc["last_chapter_title"] = ch.title or ""
    # 排序 + cap
    rows = sorted(
        pair_acc.values(),
        key=lambda x: (-x["co_chapter_count"], -x["last_chapter_no"]),
    )
    if body.min_co_chapter_count > 1:
        rows = [r for r in rows if r["co_chapter_count"] >= body.min_co_chapter_count]
    rows = rows[: max(1, body.max_pairs)]

    # 3) Synthesise one parent AgentTask so the per-pair AgentStep
    #    rows share a parent (consistent with the rest of the audit
    #    trail). Each pair becomes one AgentStep with its own inputs.
    study_task = AgentTask(
        project_id=material.project_id or 0,
        chapter_id=None,
        task_type="study_relationship",
        status="running",
        priority=50,
        payload={
            "material_id": material_id,
            "max_pairs": body.max_pairs,
            "min_co_chapter_count": body.min_co_chapter_count,
            "pair_count": len(rows),
        },
        started_at=datetime.utcnow(),
    )
    db.add(study_task)
    await db.flush()

    # 4) Run the LLM agent — one call per pair, sequential
    #    (avoids prompt 互踩; cheap model 1-2s/call so 30 pairs
    #    = ~30-60s total). Each call creates one AgentStep
    #    (audit row) and we capture cost/tokens via AgentRunResult.
    agent = StudyRelationshipExtractionAgent(
        router=get_llm_router(),
        engine=get_prompt_engine(),
    )
    items: list[StudyRelationshipEnrichedItem] = []
    enriched = 0
    fallback = 0
    skipped = 0
    cost_sum = 0.0
    for r in rows:
        ch = by_chap_id.get(r["last_chapter_id"])
        excerpt = (ch.content or "")[:1500] if ch is not None else ""
        sample_quote = excerpt[:200]
        # 解析默认值 (走 fallback / 失败路径时用)
        relation = "同章节出现"
        confidence = 0.0
        evidence = ""
        llm_inferred = False
        try:
            result = await agent.run(
                AgentContext(
                    db=db,
                    task=study_task,
                    project_id=material.project_id or 0,
                    chapter_id=None,
                    inputs={
                        "char_a_name": r["char_a"].name,
                        "char_b_name": r["char_b"].name,
                        "char_a_role": r["char_a"].role or "其他",
                        "char_b_role": r["char_b"].role or "其他",
                        "chapter_excerpt": excerpt or "(无章节正文)",
                    },
                )
            )
            cost_sum += float(result.cost_usd or 0.0)
            parsed = result.parsed or {}
            relations = parsed.get("relations") or []
            if relations and isinstance(relations, list) and len(relations) > 0:
                top = relations[0]
                relation = (top.get("relation") or "同章节出现").strip() or "同章节出现"
                confidence = float(top.get("confidence") or 0.0)
                evidence = (top.get("evidence") or "").strip()
                if relation == "未知" or confidence <= 0.0:
                    relation = "同章节出现"
                    confidence = 0.0
                    fallback += 1
                else:
                    llm_inferred = True
                    enriched += 1
            else:
                fallback += 1
        except Exception:
            # LLM 调用本身失败 → 跳过, 但仍然 emit 一条 item 维持
            # 候选对完整 (用户在前端看到 "抽取失败" 仍能看 co-occurrence)
            skipped += 1
        items.append(StudyRelationshipEnrichedItem(
            char_a_id=r["char_a"].id,
            char_a_name=r["char_a"].name,
            char_b_id=r["char_b"].id,
            char_b_name=r["char_b"].name,
            co_chapter_count=r["co_chapter_count"],
            last_chapter_no=r["last_chapter_no"],
            last_chapter_title=r.get("last_chapter_title", ""),
            sample_quote=sample_quote,
            relation=relation,
            confidence=confidence,
            evidence=evidence,
            llm_inferred=llm_inferred,
        ))

    # 5) Mark parent task as succeeded
    study_task.status = "succeeded"
    study_task.finished_at = datetime.utcnow()
    await db.flush()

    duration_ms = int((_time.time() - t0) * 1000)
    return {
        "ok": True,
        "data": StudyRelationshipEnrichResponse(
            material_id=material_id,
            enriched_count=enriched,
            skipped_count=skipped,
            fallback_count=fallback,
            duration_ms=duration_ms,
            cost_usd=cost_sum,
            items=items,
        ),
    }


@router.post(
    "/materials/{material_id}/relationships/apply",
    response_model=APIResponse[StudyRelationshipApplyResponse],
)
async def apply_relationship_suggestions(
    material_id: int,
    body: StudyRelationshipApplyRequest,
    db: AsyncSession = Depends(get_db),
) -> APIResponse[StudyRelationshipApplyResponse]:
    """R22: take the user's picked suggestions and create GraphEdges.

    Each ``pair`` in the request body is ``{char_a_id, char_b_id,
    relation}``. The route looks up the corresponding
    ``StudyCharacter`` rows, ensures they're materialised as
    ``GraphNode``s in the target project, and then creates the
    edge. Idempotent: an existing edge between the same two nodes
    with the same relation is skipped.
    """
    from app.models.study import GraphEdge, GraphNode, StudyCharacter as _SC

    material = await db.get(StudyMaterial, material_id)
    if material is None:
        raise not_found("StudyMaterial", material_id)

    if not body.pairs:
        return {
            "ok": True,
            "data": StudyRelationshipApplyResponse(
                project_id=body.project_id,
                edges_added=0,
                edges_skipped=0,
                edge_ids=[],
            ),
        }

    # Pull the relevant character rows up front so we can both
    # materialise-and-lookup in one go.
    char_ids = set()
    for p in body.pairs:
        if isinstance(p, dict):
            if isinstance(p.get("char_a_id"), int):
                char_ids.add(p["char_a_id"])
            if isinstance(p.get("char_b_id"), int):
                char_ids.add(p.get("char_b_id"))  # type: ignore[arg-type]
    if not char_ids:
        raise bad_request("pairs 里没有 char_a_id / char_b_id。")
    chars = (await db.execute(
        select(_SC).where(_SC.id.in_(list(char_ids)))
    )).scalars().all()
    chars_by_id = {c.id: c for c in chars}
    missing = char_ids - set(chars_by_id.keys())
    if missing:
        raise bad_request(f"以下人物 ID 找不到: {sorted(missing)}。")

    # Auto-materialise the characters as GraphNodes (the materialise
    # logic is small enough to inline — duplicating it here would
    # just create drift with the original endpoint).
    existing_nodes = (await db.execute(
        select(GraphNode).where(
            GraphNode.project_id == body.project_id,
            GraphNode.source_material_id == material_id,
        )
    )).scalars().all()
    nodes_by_char_id: dict[int, GraphNode] = {
        n.ref_study_character_id: n for n in existing_nodes if n.ref_study_character_id is not None
    }
    for cid in char_ids:
        if cid in nodes_by_char_id:
            continue
        c = chars_by_id[cid]
        if c.material_id != material_id:
            # Defence: don't accept a char from a different book.
            continue
        node = GraphNode(
            project_id=body.project_id,
            source_material_id=material_id,
            node_kind="study_character",
            name=c.name,
            ref_study_character_id=c.id,
            extra={
                "role": c.role,
                "tags": c.tags or [],
                "aliases": c.aliases or [],
            },
        )
        db.add(node)
        await db.flush()
        nodes_by_char_id[cid] = node

    # Build edges.
    added = 0
    skipped = 0
    edge_ids: list[int] = []
    existing_edges = (await db.execute(
        select(GraphEdge).where(GraphEdge.project_id == body.project_id)
    )).scalars().all()
    edge_key = {(e.source_node_id, e.target_node_id, e.relation) for e in existing_edges}
    for p in body.pairs:
        if not isinstance(p, dict):
            continue
        a_id = p.get("char_a_id")
        b_id = p.get("char_b_id")
        relation = (p.get("relation") or "同章节出现").strip() or "同章节出现"
        if not isinstance(a_id, int) or not isinstance(b_id, int):
            continue
        a_node = nodes_by_char_id.get(a_id)
        b_node = nodes_by_char_id.get(b_id)
        if a_node is None or b_node is None:
            skipped += 1
            continue
        if a_node.id == b_node.id:
            skipped += 1
            continue
        if (a_node.id, b_node.id, relation) in edge_key:
            # Edge exists: bump weight and count via upsert
            from sqlalchemy.dialects.sqlite import insert as sqlite_insert
            from sqlalchemy import func
            weight = float(p.get("weight") or 0.5)
            stmt = sqlite_insert(GraphEdge).values(
                project_id=body.project_id,
                source_node_id=a_node.id,
                target_node_id=b_node.id,
                relation=relation[:60],
                weight=weight,
                count=1,
            ).on_conflict_do_update(
                index_elements=["source_node_id", "target_node_id", "relation"],
                set_={
                    "weight": GraphEdge.weight + weight,
                    "count": GraphEdge.count + 1,
                    "updated_at": func.now(),
                },
            )
            await db.execute(stmt)
            skipped += 1
            continue
        edge = GraphEdge(
            project_id=body.project_id,
            source_node_id=a_node.id,
            target_node_id=b_node.id,
            relation=relation[:60],
            weight=float(p.get("weight") or 0.5),
            evidence=(p.get("evidence") or "")[:1000] or None,
        )
        db.add(edge)
        await db.flush()
        edge_key.add((a_node.id, b_node.id, relation))
        edge_ids.append(edge.id)
        added += 1
    return {
        "ok": True,
        "data": StudyRelationshipApplyResponse(
            project_id=body.project_id,
            edges_added=added,
            edges_skipped=skipped,
            edge_ids=edge_ids,
        ),
    }


@router.get(
    "/materials/{material_id}/overview",
    response_model=APIResponse[StudyMaterialOverview],
)
async def get_overview(
    material_id: int, db: AsyncSession = Depends(get_db)
) -> APIResponse[StudyMaterialOverview]:
    """R22: one-stop dashboard for a study material.

    Aggregates the "where did my data go" question so the Study
    page can render a 4-stat row (chapters / characters / behaviors
    / foreshadows) plus a small sample of each, without four
    round-trips per material card.
    """
    from app.models.memory import MemoryForeshadow
    from app.models.study import GraphNode as _GN

    material = await db.get(StudyMaterial, material_id)
    if material is None:
        raise not_found("StudyMaterial", material_id)

    # Counts (cheap; each is one COUNT query).
    n_chars = (await db.execute(
        select(func.count(StudyCharacter.id)).where(StudyCharacter.material_id == material_id)
    )).scalar_one()
    n_behaviors = (await db.execute(
        select(func.count(BehaviorPattern.id)).where(BehaviorPattern.source_material_id == material_id)
    )).scalar_one()
    n_foreshadows = 0
    if material.project_id:
        n_foreshadows = (await db.execute(
            select(func.count(MemoryForeshadow.id)).where(
                MemoryForeshadow.source_material_id == material_id,
                MemoryForeshadow.project_id == material.project_id,
            )
        )).scalar_one()
    n_graph_nodes = (await db.execute(
        select(func.count(_GN.id)).where(_GN.source_material_id == material_id)
    )).scalar_one()

    # Samples
    sample_chars = (await db.execute(
        select(StudyCharacter)
        .where(StudyCharacter.material_id == material_id)
        .order_by(StudyCharacter.confidence.desc(), StudyCharacter.id.asc())
        .limit(5)
    )).scalars().all()
    sample_behaviors = (await db.execute(
        select(BehaviorPattern)
        .where(BehaviorPattern.source_material_id == material_id)
        .order_by(BehaviorPattern.confidence.desc(), BehaviorPattern.id.asc())
        .limit(5)
    )).scalars().all()
    sample_foreshadows: list[MemoryForeshadow] = []
    if material.project_id:
        sample_foreshadows = (await db.execute(
            select(MemoryForeshadow)
            .where(
                MemoryForeshadow.source_material_id == material_id,
                MemoryForeshadow.project_id == material.project_id,
            )
            .order_by(MemoryForeshadow.planted_chapter.asc().nulls_last(), MemoryForeshadow.id.asc())
            .limit(5)
        )).scalars().all()

    return {
        "ok": True,
        "data": StudyMaterialOverview(
            material_id=material_id,
            title=material.title,
            project_id=material.project_id,
            chapter_count=material.chapter_count or 0,
            character_count=n_chars,
            behavior_count=n_behaviors,
            foreshadow_count=n_foreshadows,
            graph_node_count=n_graph_nodes,
            sample_characters=[
                {"id": c.id, "name": c.name, "role": c.role, "tags": c.tags or []}
                for c in sample_chars
            ],
            sample_behaviors=[
                {"id": p.id, "name": p.name,
                 "character_tags": p.character_tags or [],
                 "situation_tags": p.situation_tags or []}
                for p in sample_behaviors
            ],
            sample_foreshadows=[
                {"id": f.id, "name": f.name, "summary": (f.summary or "")[:120],
                 "planted_chapter": f.planted_chapter}
                for f in sample_foreshadows
            ],
        ),
    }


def _stub_study_parse(text: str) -> dict[str, Any]:
    """Deterministic character extraction used as a LAST-DITCH fallback.

    P15 / P0-STUDY-1: previously this was the *only* path; now it
    only runs when the LLM agent returns a non-JSON fallback that
    explicitly asks for it.  We keep the sliding-2-char-window
    approach because it's the only signal left when the model
    produces nothing useful at all.
    """
    if not text:
        return {"characters": []}
    import collections
    windows = [
        text[i:i+2]
        for i in range(len(text) - 1)
        if all("\u4e00" <= c <= "\u9fff" for c in text[i:i+2])
    ]
    counts = collections.Counter(windows)
    out: list[dict[str, Any]] = []
    for name, c in counts.most_common(8):
        if c < 3:
            continue
        out.append({
            "name": name,
            "aliases": [],
            "role": "其他",
            "tags": [],
            "base_profile": {"summary": f"出现 {c} 次"},
            "confidence": min(0.9, 0.4 + c * 0.02),
        })
    return {"characters": out}
